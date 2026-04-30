"""生成完整刑事案件测试用例集。

模拟真实案件：李某涉嫌故意伤害案
包含：立案决定书、拘留证、逮捕证、多份审讯笔录、多份询问笔录、
      法医鉴定、现场勘查笔录、判决书、看守所登记表、证据清单等。
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
from openpyxl import Workbook

CASE_DIR = os.path.join(os.path.dirname(__file__), "cases", "005_li_ Assault_case")
os.makedirs(CASE_DIR, exist_ok=True)


def set_doc_style(doc):
    """设置Word文档默认样式。"""
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_title(doc, text, size=16, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def add_heading_text(doc, text, size=14):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = True
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def add_line(doc, text, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def add_blank(doc, n=1):
    for _ in range(n):
        doc.add_paragraph()


def add_table_with_data(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


# ============================================================
# 文档1：立案决定书
# ============================================================
doc = Document()
set_doc_style(doc)
add_title(doc, "北京市公安局海淀分局")
add_title(doc, "立案决定书", size=18)
add_blank(doc)
add_line(doc, "京海公立字〔2024〕第0567号", size=11)
add_blank(doc)
add_line(doc, "案件来源：110报警指挥中心转接")
add_line(doc, "报案时间：2024年3月15日23时45分")
add_line(doc, "报案人：王某某（被害人）")
add_blank(doc)
add_line(doc, "经初步审查，认为符合立案条件，根据《中华人民共和国刑事诉讼法》第一百一十二条之规定，决定对以下案件立案侦查：")
add_blank(doc)
add_line(doc, "案件名称：李某涉嫌故意伤害案")
add_line(doc, "案件编号：BJHD-2024-刑立字第0567号")
add_line(doc, "案件类型：刑事案件 - 故意伤害")
add_line(doc, "犯罪嫌疑人：李某，男，1988年6月12日出生，汉族，河北省石家庄市人")
add_line(doc, "身份证号：13010219880612XXXX")
add_line(doc, "住址：北京市海淀区中关村大街88号院3号楼5层502室")
add_line(doc, "简要案情：2024年3月15日22时30分许，犯罪嫌疑人李某在海淀区中关村大街'蓝旗营'烧烤店门口，因停车纠纷与被害人王某某发生口角，继而升级为肢体冲突。李某持酒瓶击打被害人头部，致被害人颅骨骨折、脑震荡，经鉴定为轻伤二级。")
add_blank(doc)
add_line(doc, "立案依据：《中华人民共和国刑事诉讼法》第一百一十二条")
add_line(doc, "立案时间：2024年3月16日0时15分")
add_blank(doc)
add_line(doc, "承办警官：张建国（警号：012345）")
add_line(doc, "审核领导：刘志刚（副所长）")
add_blank(doc, 3)
add_line(doc, "北京市公安局海淀分局（盖章）")
add_line(doc, "2024年3月16日")
doc.save(os.path.join(CASE_DIR, "01_立案决定书.docx"))
print("✓ 01_立案决定书.docx")

# ============================================================
# 文档2：拘留证
# ============================================================
doc = Document()
set_doc_style(doc)
add_title(doc, "北京市公安局海淀分局")
add_title(doc, "拘留证", size=18)
add_blank(doc)
add_line(doc, "京海拘字〔2024〕第0234号")
add_blank(doc)
add_line(doc, "被拘留人：李某")
add_line(doc, "性别：男  出生日期：1988年6月12日")
add_line(doc, "民族：汉族  籍贯：河北省石家庄市")
add_line(doc, "身份证号：13010219880612XXXX")
add_line(doc, "住址：北京市海淀区中关村大街88号院3号楼5层502室")
add_line(doc, "工作单位：北京某科技有限公司  职务：软件工程师")
add_blank(doc)
add_line(doc, "经查明：李某于2024年3月15日22时30分许，在海淀区中关村大街'蓝旗营'烧烤店门口，因停车纠纷与被害人王某某发生口角，继而升级为肢体冲突。李某持酒瓶击打被害人头部，致被害人颅骨骨折、脑震荡，经鉴定为轻伤二级。")
add_blank(doc)
add_line(doc, "根据《中华人民共和国刑事诉讼法》第八十二条之规定，决定对李某予以刑事拘留。")
add_blank(doc)
add_line(doc, "拘留地点：北京市海淀区看守所")
add_line(doc, "拘留时间：2024年3月17日10时00分")
add_blank(doc, 2)
add_line(doc, "批准人：刘志刚（副局长）")
add_blank(doc, 2)
add_line(doc, "执行警官：张建国（警号：012345）、赵强（警号：012346）")
doc.save(os.path.join(CASE_DIR, "02_拘留证.docx"))
print("✓ 02_拘留证.docx")

# ============================================================
# 文档3：逮捕证
# ============================================================
doc = Document()
set_doc_style(doc)
add_title(doc, "北京市公安局海淀分局")
add_title(doc, "逮捕证", size=18)
add_blank(doc)
add_line(doc, "京海捕字〔2024〕第0189号")
add_blank(doc)
add_line(doc, "被逮捕人：李某")
add_line(doc, "性别：男  出生日期：1988年6月12日")
add_line(doc, "涉嫌罪名：故意伤害罪")
add_blank(doc)
add_line(doc, "经本局侦查查明，犯罪嫌疑人李某涉嫌故意伤害罪一案，经北京市海淀区人民检察院批准（京海检批捕字〔2024〕第0145号），决定对李某予以逮捕。")
add_blank(doc)
add_line(doc, "逮捕时间：2024年4月2日14时30分")
add_line(doc, "逮捕地点：北京市海淀区看守所")
add_blank(doc, 2)
add_line(doc, "批准机关：北京市海淀区人民检察院")
add_line(doc, "批准文号：京海检批捕字〔2024〕第0145号")
add_line(doc, "批准日期：2024年3月31日")
doc.save(os.path.join(CASE_DIR, "03_逮捕证.docx"))
print("✓ 03_逮捕证.docx")

# ============================================================
# 文档4：第一次审讯笔录
# ============================================================
doc = Document()
set_doc_style(doc)
add_title(doc, "北京市海淀区公安分局", size=14)
add_title(doc, "审讯笔录", size=18)
add_blank(doc)
add_line(doc, "时间：2024年3月17日10时30分至2024年3月17日12时15分")
add_line(doc, "地点：北京市海淀区公安局海淀派出所审讯室")
add_line(doc, "审讯人：张建国（警号012345）、赵强（警号012346）")
add_line(doc, "记录人：赵强")
add_line(doc, "被审讯人：李某")
add_line(doc, "案由：涉嫌故意伤害罪")
add_blank(doc)
add_heading_text(doc, "【权利义务告知】")
add_line(doc, "问：李某，我们是北京市公安局海淀分局的民警，这是我们的警官证。")
add_line(doc, "答：好的，我看到了。")
add_line(doc, "问：根据《中华人民共和国刑事诉讼法》的规定，你有权委托辩护人。如果你没有委托辩护人，可以申请法律援助。你听清楚了吗？")
add_line(doc, "答：听清楚了。我还没有委托辩护人。")
add_line(doc, "问：你有权申请回避，即如果你认为我们与本案有利害关系，有权要求我们回避。你是否申请回避？")
add_line(doc, "答：不申请。")
add_line(doc, "问：你必须如实回答我们的问题，不得作虚假陈述，也不得隐匿证据。你听清楚了吗？")
add_line(doc, "答：听清楚了。")
add_blank(doc)
add_heading_text(doc, "【案情讯问】")
add_line(doc, "问：李某，你知道为什么把你带到派出所吗？")
add_line(doc, "答：知道，是因为3月15日晚上我跟那个人打了一架。")
add_line(doc, "问：你详细说一下当时的情况。")
add_line(doc, "答：3月15日晚上大概10点半左右，我骑电动车到蓝旗营烧烤店门口准备停车，把车停在路边。结果一个男的从烧烤店出来，开一辆黑色的奔驰，我的电动车挡住了他倒车的路。他就下车让我挪车。")
add_line(doc, "问：那个人有什么特征？")
add_line(doc, "答：大约40多岁，身高1米75左右，比较壮，穿着深色夹克。")
add_line(doc, "问：然后发生了什么？")
add_line(doc, "答：他说'你怎么停的，挡着路了'，态度很不好。我说'我等一下就走'。他就不干了，过来推我。他力气很大，把我推了个趔趄。我也还手了，推了他一下。然后他拿起地上的一个酒瓶就打我头，我头上挨了一下，很疼。我当时喝了不少酒，脑子一热，就也拿了一个酒瓶砸他头上了。")
add_line(doc, "问：你当天晚上饮酒了吗？饮了多少？")
add_line(doc, "答：喝了，大概喝了四五瓶啤酒。")
add_line(doc, "问：你拿酒瓶打了他几下？")
add_line(doc, "答：就一下。我没想多打他，就是一下。")
add_line(doc, "问：打到他什么部位？")
add_line(doc, "答：头部，左侧。")
add_line(doc, "问：他被打之后什么反应？")
add_line(doc, "答：他当时就倒了，头部流血，我就赶紧打120叫救护车了。后来我也打了110报警，说有人受伤了。")
add_line(doc, "问：你在现场等候公安人员了吗？")
add_line(doc, "答：是的，我一直在现场等着，没有离开。")
add_line(doc, "问：公安人员到了之后你有反抗或者逃跑吗？")
add_line(doc, "答：没有，我主动配合调查的。")
add_blank(doc)
add_heading_text(doc, "【最后陈述】")
add_line(doc, "问：你还有什么要补充的吗？")
add_line(doc, "答：我确实不对，当时喝多了，一时冲动。我愿意赔偿医药费。")
add_line(doc, "问：以上笔录你看过吗？和你说的相符吗？")
add_line(doc, "答：看过了，和我说的相符。")
add_blank(doc, 2)
add_line(doc, "被审讯人签名：李某（捺印）")
add_line(doc, "审讯人签名：张建国 赵强")
add_line(doc, "记录人签名：赵强")
doc.save(os.path.join(CASE_DIR, "04_第一次审讯笔录_20240317.docx"))
print("✓ 04_第一次审讯笔录_20240317.docx")

# ============================================================
# 文档5：第二次审讯笔录
# ============================================================
doc = Document()
set_doc_style(doc)
add_title(doc, "北京市海淀区公安分局", size=14)
add_title(doc, "审讯笔录（第二次）", size=18)
add_blank(doc)
add_line(doc, "时间：2024年3月22日14时00分至2024年3月22日15时30分")
add_line(doc, "地点：北京市海淀区看守所审讯室")
add_line(doc, "审讯人：张建国（警号012345）")
add_line(doc, "记录人：张建国")
add_line(doc, "被审讯人：李某")
add_blank(doc)
add_line(doc, "问：李某，今天继续讯问你关于3月15日故意伤害一案。你先说一下当时使用的酒瓶的情况。")
add_line(doc, "答：当时烧烤店门口地上有几个空酒瓶，我随手拿了一个，就是平时啤酒瓶那种。")
add_line(doc, "问：那个酒瓶是你先拿起来使用的，还是对方先使用的？")
add_line(doc, "答：是他先拿了一个酒瓶打我头，我被打之后才从地上拿了一个反击的。")
add_line(doc, "问：你头部的伤怎么处理了？")
add_line(doc, "答：事后去医院缝了三针。有医院的诊断证明。")
add_line(doc, "问：你与被害人王某某之前认识吗？")
add_line(doc, "答：不认识，完全不认识。")
add_line(doc, "问：你以前有没有打架斗殴的前科？")
add_line(doc, "答：没有，我是第一次。")
add_line(doc, "问：你在北京某科技有限公司做什么工作？")
add_line(doc, "答：我是软件工程师，负责后端开发，月薪大概两万。")
add_line(doc, "问：案发当天你为什么去蓝旗营烧烤店？")
add_line(doc, "答：公司同事聚餐，我吃完饭先走了一会儿。")
add_line(doc, "问：你有没有什么对你有利的证据或证人？")
add_line(doc, "答：烧烤店门口有监控摄像头，应该能拍到全过程。还有几个同事可以证明我在哪里喝过酒。烧烤店老板也可以证明是对方先动的手。")
add_blank(doc, 2)
add_line(doc, "被审讯人签名：李某（捺印）")
add_line(doc, "审讯人签名：张建国")
doc.save(os.path.join(CASE_DIR, "05_第二次审讯笔录_20240322.docx"))
print("✓ 05_第二次审讯笔录_20240322.docx")

# ============================================================
# 文档6：被害人询问笔录
# ============================================================
doc = Document()
set_doc_style(doc)
add_title(doc, "北京市海淀区公安分局", size=14)
add_title(doc, "询问笔录（被害人）", size=18)
add_blank(doc)
add_line(doc, "时间：2024年3月16日09时00分至2024年3月16日10时30分")
add_line(doc, "地点：北京大学人民医院病房")
add_line(doc, "询问人：张建国（警号012345）")
add_line(doc, "记录人：赵强")
add_line(doc, "被询问人：王某某（被害人）")
add_line(doc, "性别：男  年龄：42岁  职业：某贸易公司销售经理")
add_blank(doc)
add_line(doc, "问：王某某，请你详细讲述一下案发经过。")
add_line(doc, "答：3月15日晚上我和几个朋友在蓝旗营烧烤店吃饭，大概10点半左右我准备走。我开的是黑色奔驰，停在店门口的停车位。我一倒车，发现后面有一辆电动车挡住了我的路。我就下车让那个骑电动车的人挪一下。")
add_line(doc, "问：对方什么态度？")
add_line(doc, "答：他说等一下就走，但一直不动。我有点生气，就说了他两句。然后他态度也不好，两人就开始吵了。")
add_line(doc, "问：争吵之后发生了什么？")
add_line(doc, "答：我推了他一下，他也有还手。然后我看到地上有个空酒瓶，我拿起来想吓唬他，就举了一下。结果他从我手上夺过去酒瓶，反手就砸我头上了。我头上当时就流血了，然后我就倒了。")
add_line(doc, "问：你头部被打了几处？")
add_line(doc, "答：就左侧一下，但力度很大。")
add_line(doc, "问：对方有没有继续攻击你？")
add_line(doc, "答：没有，砸完我就站在那儿了，后来他打了120和110。")
add_line(doc, "问：你平时和对方认识吗？")
add_line(doc, "答：不认识，第一次见。")
add_line(doc, "问：你当天晚上饮酒了吗？")
add_line(doc, "答：喝了，喝了大概半斤白酒。")
add_blank(doc, 2)
add_line(doc, "被询问人签名：王某某（捺印）")
add_line(doc, "询问人签名：张建国")
doc.save(os.path.join(CASE_DIR, "06_被害人询问笔录_20240316.docx"))
print("✓ 06_被害人询问笔录_20240316.docx")

# ============================================================
# 文档7：证人证言（烧烤店老板）
# ============================================================
doc = Document()
set_doc_style(doc)
add_title(doc, "北京市海淀区公安分局", size=14)
add_title(doc, "询问笔录（证人）", size=18)
add_blank(doc)
add_line(doc, "时间：2024年3月16日14时00分至2024年3月16日15时00分")
add_line(doc, "地点：蓝旗营烧烤店办公室")
add_line(doc, "询问人：张建国（警号012345）")
add_line(doc, "记录人：赵强")
add_line(doc, "被询问人：陈某某（证人）")
add_line(doc, "身份：蓝旗营烧烤店老板")
add_blank(doc)
add_line(doc, "问：陈某某，请你如实陈述3月15日晚在烧烤店门口发生的冲突情况。")
add_line(doc, "答：那天晚上大概10点半，我当时在店门口抽烟。看到门口有两个人吵架。一个开黑色奔驰的男的（后来知道是被害人王某某）和一个骑电动车的年轻男的（后来知道是嫌疑人李某）。王某某的车被电动车挡住了，就让李某挪车。两个人就吵起来了。")
add_line(doc, "问：双方吵架的内容你听到了什么？")
add_line(doc, "答：王某某说'你怎么停车的，挡着路了'，李某说'我等一下就走'。然后王某某语气就比较冲了，说了几句不太好听的。李某也不示弱。")
add_line(doc, "问：后来有没有发生肢体冲突？")
add_line(doc, "答：有。王某某先推了李某一下，李某也还手推了王某某一下。然后王某某从地上拿了一个酒瓶举起来，李某就从他手里夺过来酒瓶，反手就砸王某某头上了。")
add_line(doc, "问：谁先拿的酒瓶？")
add_line(doc, "答：是王某某先拿的，他从地上捡了一个酒瓶举起来，然后李某夺过去了。")
add_line(doc, "问：你看到李某打了几下？")
add_line(doc, "答：就一下。砸完王某某就倒了。")
add_line(doc, "问：李某打人之后有没有逃跑？")
add_line(doc, "答：没有，他一直在现场，还打了120和110。")
add_line(doc, "问：你们店门口的监控能拍到整个过程吗？")
add_line(doc, "答：可以的，门口有两个摄像头，应该能拍到。")
add_blank(doc, 2)
add_line(doc, "被询问人签名：陈某某（捺印）")
add_line(doc, "询问人签名：张建国")
doc.save(os.path.join(CASE_DIR, "07_证人证言_烧烤店老板_20240316.docx"))
print("✓ 07_证人证言_烧烤店老板_20240316.docx")

# ============================================================
# 文档8：证人证言（路过市民）
# ============================================================
doc = Document()
set_doc_style(doc)
add_title(doc, "北京市海淀区公安分局", size=14)
add_title(doc, "询问笔录（证人）", size=18)
add_blank(doc)
add_line(doc, "时间：2024年3月16日10时00分至2024年3月16日10时40分")
add_line(doc, "地点：中关村大街派出所")
add_line(doc, "询问人：赵强（警号012346）")
add_line(doc, "记录人：赵强")
add_line(doc, "被询问人：刘某（证人）")
add_line(doc, "身份：路过市民")
add_blank(doc)
add_line(doc, "问：刘某，请你讲述一下3月15日晚上看到的情况。")
add_line(doc, "答：我那天晚上从烧烤店门口路过，看到两个人在打架。一个一个高壮的人推了一个瘦一点的年轻一下，然后两个人就扭在一起了。后来我看到瘦的那个拿了个什么东西打了高壮的那个人头部一下，高壮的人就倒了。")
add_line(doc, "问：你看清打人的工具了吗？")
add_line(doc, "答：看不太清，像是个瓶子之类的东西。")
add_line(doc, "问：打人之后打人的人有没有逃跑？")
add_line(doc, "答：没有，他一直在现场等着。")
add_blank(doc, 2)
add_line(doc, "被询问人签名：刘某（捺印）")
add_line(doc, "询问人签名：赵强")
doc.save(os.path.join(CASE_DIR, "08_证人证言_路过市民_20240316.docx"))
print("✓ 08_证人证言_路过市民_20240316.docx")

# ============================================================
# 文档9：法医鉴定意见书
# ============================================================
doc = Document()
set_doc_style(doc)
add_title(doc, "北京市公安司法鉴定中心", size=14)
add_title(doc, "法医学人体损伤程度鉴定意见书", size=16)
add_blank(doc)
add_line(doc, "鉴定编号：京公法临鉴字〔2024〕第0892号")
add_line(doc, "委托单位：北京市公安局海淀分局")
add_line(doc, "委托日期：2024年3月16日")
add_line(doc, "鉴定日期：2024年3月18日")
add_blank(doc)
add_heading_text(doc, "一、被鉴定人信息")
add_line(doc, "姓名：王某某  性别：男  年龄：42岁")
add_line(doc, "身份证号：11010819820115XXXX")
add_blank(doc)
add_heading_text(doc, "二、案情摘要")
add_line(doc, "2024年3月15日22时30分许，被鉴定人王某某在海淀区中关村大街'蓝旗营'烧烤店门口与他人发生冲突，被人用酒瓶击打头部，致伤。")
add_blank(doc)
add_heading_text(doc, "三、检验情况")
add_line(doc, "1. 头部：左侧颞部可见一4.5厘米弧形创口，创缘不整，创角一钝一锐，深达皮下及颅骨。")
add_line(doc, "2. CT检查：左侧颞骨线性骨折，骨折线长约3.2厘米。")
add_line(doc, "3. 脑震荡症状：伤后有短暂意识丧失（约3分钟），伴有恶心、呕吐。")
add_line(doc, "4. 其他：面部软组织挫伤2处，右前臂擦伤1处。")
add_blank(doc)
add_heading_text(doc, "四、鉴定意见")
add_line(doc, "被鉴定人王某某头部损伤程度为轻伤二级。")
add_line(doc, "依据：《人体损伤程度鉴定标准》第5.1.4条——颅骨骨折，评定为轻伤二级。")
add_blank(doc, 2)
add_line(doc, "鉴定人：孙伟明（主检法医师）")
add_line(doc, "审核人：张文华（副主任法医师）")
add_line(doc, "北京市公安司法鉴定中心（盖章）")
add_line(doc, "2024年3月18日")
doc.save(os.path.join(CASE_DIR, "09_法医鉴定意见书_20240318.docx"))
print("✓ 09_法医鉴定意见书_20240318.docx")

# ============================================================
# 文档10：现场勘验检查笔录
# ============================================================
doc = Document()
set_doc_style(doc)
add_title(doc, "北京市公安局海淀分局", size=14)
add_title(doc, "现场勘验检查笔录", size=18)
add_blank(doc)
add_line(doc, "勘验编号：京海勘字〔2024〕第0445号")
add_line(doc, "勘验时间：2024年3月16日08时00分至2024年3月16日10时00分")
add_line(doc, "勘验地点：海淀区中关村大街88号蓝旗营烧烤店门口")
add_line(doc, "勘验人：李明（痕检工程师）、王强（法医）")
add_blank(doc)
add_heading_text(doc, "一、现场环境")
add_line(doc, "现场位于海淀区中关村大街88号蓝旗营烧烤店门前人行道与临时停车区域交汇处。烧烤店门口设有2个临时停车位，地面为水泥硬化路面。")
add_blank(doc)
add_heading_text(doc, "二、勘验情况")
add_line(doc, "1. 烧烤店入口上方安装监控摄像头2个，已提取视频数据。")
add_line(doc, "2. 现场地面发现血迹一处（已提取样本，经鉴定与被害人王某某血型一致）。")
add_line(doc, "3. 现场提取玻璃碎片若干（经鉴定为啤酒瓶碎片，与作案工具特征相符）。")
add_line(doc, "4. 现场地面发现搏斗痕迹（擦痕、脚印等）。")
add_blank(doc)
add_heading_text(doc, "三、提取物品")
add_line(doc, "1. 血迹棉签2支（编号：XJ-001、XJ-002）")
add_line(doc, "2. 玻璃碎片3块（编号：BL-001至BL-003）")
add_line(doc, "3. 现场照片12张")
add_blank(doc, 2)
add_line(doc, "勘验人签名：李明 王强")
add_line(doc, "见证人签名：陈某某（烧烤店老板）")
doc.save(os.path.join(CASE_DIR, "10_现场勘验笔录_20240316.docx"))
print("✓ 10_现场勘验笔录_20240316.docx")

# ============================================================
# 文档11：监控视频分析说明
# ============================================================
doc = Document()
set_doc_style(doc)
add_title(doc, "北京市公安局海淀分局", size=14)
add_title(doc, "监控视频分析说明", size=18)
add_blank(doc)
add_line(doc, "视频来源：蓝旗营烧烤店门口监控摄像头（2个）")
add_line(doc, "视频时间段：2024年3月15日22时20分至22时40分")
add_line(doc, "分析人：张建国")
add_blank(doc)
add_heading_text(doc, "视频关键帧时间戳")
add_line(doc, "22:25:13 - 嫌疑人李某骑电动车到达现场，停放在临时停车位")
add_line(doc, "22:28:45 - 被害人王某某从烧烤店出来，走向黑色奔驰车")
add_line(doc, "22:29:10 - 王某某与李某开始交谈（口角）")
add_line(doc, "22:29:45 - 王某某推了李某一下")
add_line(doc, "22:30:02 - 李某还手推了王某某一下")
add_line(doc, "22:30:15 - 王某某弯腰从地上拿起一个酒瓶")
add_line(doc, "22:30:22 - 李某从王某某手中夺过酒瓶")
add_line(doc, "22:30:28 - 李某持酒瓶击打王某某头部")
add_line(doc, "22:30:32 - 王某某倒地")
add_line(doc, "22:30:50 - 李某拨打手机（120急救电话）")
add_line(doc, "22:31:15 - 李某再次拨打手机（110报警电话）")
add_line(doc, "22:32:00 - 李某在现场等候，未有离开动作")
add_blank(doc, 2)
add_line(doc, "分析人签名：张建国")
doc.save(os.path.join(CASE_DIR, "11_监控视频分析说明_20240317.docx"))
print("✓ 11_监控视频分析说明_20240317.docx")

# ============================================================
# 文档12：刑事判决书
# ============================================================
doc = Document()
set_doc_style(doc)
add_title(doc, "北京市海淀区人民法院", size=14)
add_title(doc, "刑事判决书", size=18)
add_blank(doc)
add_line(doc, "（2024）京0108刑初字1234号")
add_blank(doc)
add_heading_text(doc, "公诉机关：北京市海淀区人民检察院")
add_heading_text(doc, "被告人：李某，男，1988年6月12日出生")
add_line(doc, "辩护人：赵某某，北京正义律师事务所律师")
add_blank(doc)
add_heading_text(doc, "案件经过")
add_line(doc, "北京市海淀区人民检察院以京海检刑诉〔2024〕第0789号起诉书指控被告人李某犯故意伤害罪，于2024年6月15日向本院提起公诉。本院依法组成合议庭，于2024年7月20日公开开庭审理了本案。现已审理终结。")
add_blank(doc)
add_heading_text(doc, "法院查明事实")
add_line(doc, "经审理查明：2024年3月15日22时30分许，被告人李某在海淀区中关村大街'蓝旗营'烧烤店门口，因停车纠纷与被害人王某某发生口角，继而升级为肢体冲突。在冲突过程中，被害人王某某先拿酒瓶欲击打被告人李某，李某夺过酒瓶后击打王某某头部，致王某某左侧颞骨骨折，经鉴定为轻伤二级。案发后，被告人李某主动拨打120和110，在现场等候公安人员处理，到案后如实供述犯罪事实。")
add_blank(doc)
add_heading_text(doc, "法院认为")
add_line(doc, "被告人李某故意伤害他人身体，致人轻伤，其行为已构成故意伤害罪，依法应予惩处。鉴于被告人李某案发后主动报警并在现场等候，到案后如实供述，具有自首情节，依法可从轻处罚。被告人李某家属已赔偿被害人的经济损失并取得谅解，可酌情从轻处罚。被害人在案件起因上存在一定过错，可酌情对被告人从轻处罚。")
add_blank(doc)
add_heading_text(doc, "判决如下")
add_line(doc, "依照《中华人民共和国刑法》第二百三十四条第一款、第六十七条第一款之规定，判决如下：")
add_blank(doc)
add_line(doc, "被告人李某犯故意伤害罪，判处有期徒刑十个月。")
add_line(doc, "（刑期从判决执行之日起计算。判决执行以前先行羁押的，羁押一日折抵刑期一日。）")
add_blank(doc, 2)
add_line(doc, "审判长：周明")
add_line(doc, "审判员：刘芳")
add_line(doc, "人民陪审员：王建国")
add_line(doc, "二〇二四年八月十日")
add_line(doc, "书记员：孙丽")
doc.save(os.path.join(CASE_DIR, "12_刑事判决书_20240810.docx"))
print("✓ 12_刑事判决书_20240810.docx")

# ============================================================
# 文档13：取保候审决定书
# ============================================================
doc = Document()
set_doc_style(doc)
add_title(doc, "北京市公安局海淀分局", size=14)
add_title(doc, "取保候审决定书", size=18)
add_blank(doc)
add_line(doc, "京海取保字〔2024〕第0089号")
add_blank(doc)
add_line(doc, "犯罪嫌疑人：李某")
add_line(doc, "涉嫌罪名：故意伤害罪")
add_blank(doc)
add_line(doc, "经本局审查认为，犯罪嫌疑人李某涉嫌故意伤害罪一案，因其具有自首情节，且无前科，不具有社会危险性，符合取保候审条件。根据《中华人民共和国刑事诉讼法》第六十七条之规定，决定对李某取保候审。")
add_blank(doc)
add_line(doc, "取保候审期限：自2024年5月10日起至2025年5月9日止")
add_line(doc, "保证金金额：人民币10,000元")
add_line(doc, "义务规定：")
add_line(doc, "1. 未经批准不得离开北京市海淀区")
add_line(doc, "2. 在传讯的时候及时到案")
add_line(doc, "3. 不得干扰证人作证")
add_line(doc, "4. 不得毁灭、伪造证据或者串供")
add_blank(doc, 2)
add_line(doc, "批准人：刘志刚（副局长）")
add_line(doc, "北京市公安局海淀分局（盖章）")
add_line(doc, "2024年5月10日")
doc.save(os.path.join(CASE_DIR, "13_取保候审决定书_20240510.docx"))
print("✓ 13_取保候审决定书_20240510.docx")

# ============================================================
# 文档14：赔偿协议书
# ============================================================
doc = Document()
set_doc_style(doc)
add_title(doc, "赔偿协议书", size=18)
add_blank(doc)
add_line(doc, "甲方（赔偿方）：李某")
add_line(doc, "乙方（受偿方）：王某某")
add_blank(doc)
add_line(doc, "鉴于：2024年3月15日，甲方因停车纠纷与乙方发生冲突，致乙方头部受伤。甲方对此深表歉意。经双方友好协商，就赔偿事宜达成如下协议：")
add_blank(doc)
add_line(doc, "一、甲方一次性赔偿乙方以下费用：")
add_line(doc, "1. 医疗费：人民币28,500元")
add_line(doc, "2. 误工费：人民币15,000元")
add_line(doc, "3. 护理费：人民币5,000元")
add_line(doc, "4. 营养费：人民币3,000元")
add_line(doc, "5. 精神损害抚慰金：人民币20,000元")
add_line(doc, "以上合计人民币71,500元")
add_blank(doc)
add_line(doc, "二、乙方收到上述赔偿款后，出具谅解书，对甲方表示谅解。")
add_line(doc, "三、本协议自双方签字之日起生效。")
add_blank(doc, 2)
add_line(doc, "甲方（签名）：李某    日期：2024年5月15日")
add_line(doc, "乙方（签名）：王某某  日期：2024年5月15日")
add_line(doc, "见证人：赵某某律师    日期：2024年5月15日")
doc.save(os.path.join(CASE_DIR, "14_赔偿协议书_20240515.docx"))
print("✓ 14_赔偿协议书_20240515.docx")

# ============================================================
# Excel 1：看守所登记表
# ============================================================
wb = Workbook()
ws = wb.active
ws.name = "看守所登记表"

headers = ["序号", "姓名", "性别", "年龄", "涉嫌罪名", "入所日期", "监室号", "健康情况", "备注"]
ws.append(headers)

data = [
    [1, "李某", "男", 36, "故意伤害", "2024-03-17", "3-12", "良好，头部缝3针", "5月10日取保"],
    [2, "张某", "男", 28, "盗窃", "2024-03-10", "2-05", "良好", ""],
    [3, "王某", "男", 45, "诈骗", "2024-03-12", "1-08", "高血压，服药中", ""],
    [4, "刘某", "男", 32, "抢劫", "2024-03-18", "4-03", "良好", ""],
    [5, "陈某", "男", 50, "故意伤害", "2024-03-20", "2-11", "糖尿病，服药中", ""],
    [6, "赵某", "女", 25, "贩毒", "2024-03-22", "5-02", "良好", ""],
    [7, "孙某", "男", 38, "危险驾驶", "2024-03-25", "1-06", "良好，已释放", "4月5日释放"],
]

for row in data:
    ws.append(row)

# ============================================================
# Excel 2：证据清单
# ============================================================
ws2 = wb.create_sheet("证据清单")

headers2 = ["序号", "证据名称", "证据类型", "来源", "提取时间", "提取人", "编号", "备注"]
ws2.append(headers2)

data2 = [
    [1, "监控视频（2个摄像头）", "视听资料", "蓝旗营烧烤店", "2024-03-16", "张建国", "SP-001", "关键证据，记录全过程"],
    [2, "玻璃碎片（3块）", "物证", "案发现场", "2024-03-16", "李明", "BL-001至003", "啤酒瓶碎片"],
    [3, "血迹样本（2份）", "物证", "案发现场", "2024-03-16", "王强", "XJ-001、002", "与被害人血型一致"],
    [4, "作案工具（酒瓶）", "物证", "案发现场", "2024-03-16", "张建国", "ZJ-001", "未找到（已碎裂）"],
    [5, "法医鉴定意见书", "鉴定意见", "北京市公安司法鉴定中心", "2024-03-18", "孙伟明", "JD-001", "轻伤二级"],
    [6, "嫌疑人第一次审讯笔录", "被告人供述", "海淀派出所", "2024-03-17", "赵强", "BL-004", "供认不讳"],
    [7, "嫌疑人第二次审讯笔录", "被告人供述", "海淀看守所", "2024-03-22", "张建国", "BL-005", "细节补充"],
    [8, "被害人询问笔录", "被害人陈述", "北京大学人民医院", "2024-03-16", "张建国", "BL-006", ""],
    [9, "证人陈某某证言", "证人证言", "蓝旗营烧烤店", "2024-03-16", "张建国", "BL-007", "证明被害人先拿酒瓶"],
    [10, "证人刘某证言", "证人证言", "中关村派出所", "2024-03-16", "赵强", "BL-008", "路过目击"],
    [11, "现场勘验笔录", "勘验笔录", "案发现场", "2024-03-16", "李明", "KY-001", ""],
    [12, "医院诊断证明", "书证", "北京大学人民医院", "2024-03-16", "张建国", "SD-001", "颅骨骨折"],
    [13, "赔偿协议书", "书证", "北京正义律师事务所", "2024-05-15", "张建国", "SD-002", "赔偿71500元"],
    [14, "谅解书", "书证", "王某某", "2024-05-15", "张建国", "SD-003", ""],
]

for row in data2:
    ws2.append(row)

# ============================================================
# Excel 3：费用清单
# ============================================================
ws3 = wb.create_sheet("费用清单")

headers3 = ["序号", "费用类型", "金额(元)", "支付方", "收款方", "日期", "凭证"]
ws3.append(headers3)

data3 = [
    [1, "医疗费", 28500, "李某", "北京大学人民医院", "2024-03-16至03-25", "医院发票"],
    [2, "误工费", 15000, "李某", "王某某", "2024-05-15", "赔偿协议"],
    [3, "护理费", 5000, "李某", "王某某", "2024-05-15", "赔偿协议"],
    [4, "营养费", 3000, "李某", "王某某", "2024-05-15", "赔偿协议"],
    [5, "精神损害抚慰金", 20000, "李某", "王某某", "2024-05-15", "赔偿协议"],
    [6, "律师费", 8000, "李某", "北京正义律师事务所", "2024-05-01", "律所发票"],
    [7, "鉴定费", 2000, "公安局", "北京市公安司法鉴定中心", "2024-03-18", "鉴定中心收据"],
    [8, "看守所费用", 1500, "李某", "海淀区看守所", "2024-03-17至05-10", "看守所收据"],
]

for row in data3:
    ws3.append(row)

wb.save(os.path.join(CASE_DIR, "15_案件登记表_看守所+证据+费用.xlsx"))
print("✓ 15_案件登记表_看守所+证据+费用.xlsx")

# Count files
import glob
files = glob.glob(os.path.join(CASE_DIR, "*"))
print(f"\n共计生成 {len(files)} 个测试文件：")
for f in sorted(files):
    print(f"  {os.path.basename(f)}")
