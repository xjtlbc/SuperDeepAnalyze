"""Convert legacy .doc files to .docx using multi-tier fallback chain.

Provides a 6-tier fallback chain:
  Tier 0: Conversion cache (file hash → cached .md)
  Tier 1: LibreOffice headless (best quality, preserves formatting)
  Tier 2: antiword + chardet encoding detection
  Tier 3: catdoc (lightweight, good Chinese support)
  Tier 4: python-docx direct attempt (in case file is actually .docx)
  Tier 5: Structured OLE extraction via olefile (last resort)
"""

import asyncio
import hashlib
import logging
import os
import shutil
import signal
import subprocess
import tempfile
from pathlib import Path

logger = logging.getLogger("app.parsing.doc_converter")

CONVERT_TIMEOUT = 120  # seconds
MAX_RETRIES = 2
RETRY_DELAY = 5  # seconds between retries

# Cache directory for converted .doc files
CACHE_DIR = Path(__file__).parent.parent.parent.parent.parent / "data" / "cache" / "doc_converted"


def _compute_file_hash(path: Path) -> str:
    """Compute SHA256 hash of file for cache key."""
    sha = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            sha.update(chunk)
    return sha.hexdigest()[:16]


def _get_cache_path(file_hash: str) -> Path:
    """Get cache file path for a given file hash."""
    CACHE_DIR.mkdir(parents=True, exist_ok=True)
    return CACHE_DIR / f"{file_hash}.md"


def _read_cache(file_hash: str) -> str | None:
    """Read cached conversion result if it exists and is non-empty."""
    cache_path = _get_cache_path(file_hash)
    if cache_path.exists():
        content = cache_path.read_text(encoding="utf-8")
        if content.strip():
            logger.info("[doc] Cache hit for hash=%s (%d chars)", file_hash, len(content))
            return content
    return None


def _write_cache(file_hash: str, content: str) -> None:
    """Write conversion result to cache."""
    cache_path = _get_cache_path(file_hash)
    cache_path.write_text(content, encoding="utf-8")
    logger.debug("[doc] Cached conversion for hash=%s", file_hash)


# ── Tool availability ──────────────────────────────────────────────────────


def is_libreoffice_available() -> bool:
    return shutil.which("soffice") is not None


def is_antiword_available() -> bool:
    return shutil.which("antiword") is not None


def is_catdoc_available() -> bool:
    return shutil.which("catdoc") is not None


# ── Tier 1: LibreOffice ────────────────────────────────────────────────────


def _kill_zombie_soffice() -> None:
    """Kill any existing zombie soffice processes before starting conversion."""
    try:
        # Send SIGTERM to all soffice.bin processes
        result = subprocess.run(
            ["pgrep", "-f", "soffice"],
            capture_output=True, text=True, timeout=5,
        )
        if result.returncode == 0 and result.stdout.strip():
            pids = result.stdout.strip().split("\n")
            logger.info("[doc] Killing %d zombie soffice process(es): %s", len(pids), ", ".join(pids))
            for pid in pids:
                try:
                    os.kill(int(pid), signal.SIGKILL)
                except (ProcessLookupError, PermissionError):
                    pass
    except Exception:
        pass  # pgrep not available or other error, continue anyway


async def convert_doc_to_docx(doc_path: Path) -> Path:
    """Convert a .doc file to .docx via LibreOffice headless with retry.

    Returns the path to the converted .docx file in a temp directory.
    Caller is responsible for cleanup.
    """
    if not is_libreoffice_available():
        raise RuntimeError(
            "LibreOffice is not installed. Cannot convert .doc files. "
            "Install with: apt-get install libreoffice-writer"
        )

    doc_path = Path(doc_path)
    if not doc_path.exists():
        raise FileNotFoundError(f"File not found: {doc_path}")

    # Kill zombie processes before starting
    _kill_zombie_soffice()

    last_error: Exception | None = None

    for attempt in range(1, MAX_RETRIES + 1):
        out_dir = tempfile.mkdtemp(prefix="doc_convert_")
        try:
            # Use --norestore to avoid LibreOffice getting stuck on recovery dialogs
            # Use env HOME to a temp dir to avoid ~/.config/libreoffice lock issues
            env = os.environ.copy()
            env["HOME"] = out_dir

            proc = await asyncio.create_subprocess_exec(
                "soffice",
                "--headless",
                "--norestore",
                "--convert-to", "docx",
                "--outdir", out_dir,
                str(doc_path),
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
                env=env,
            )

            try:
                stdout, stderr = await asyncio.wait_for(
                    proc.communicate(), timeout=CONVERT_TIMEOUT
                )
            except asyncio.TimeoutError:
                # Kill the entire process group, not just the main process
                try:
                    os.killpg(os.getpgid(proc.pid), signal.SIGKILL)
                except (ProcessLookupError, OSError):
                    proc.kill()
                last_error = RuntimeError(
                    f"LibreOffice conversion timed out after {CONVERT_TIMEOUT}s "
                    f"for {doc_path.name} (attempt {attempt}/{MAX_RETRIES})"
                )
                logger.warning(str(last_error))
                _cleanup_temp(out_dir)
                if attempt < MAX_RETRIES:
                    await asyncio.sleep(RETRY_DELAY)
                continue

            if proc.returncode != 0:
                err = stderr.decode(errors="replace").strip()
                last_error = RuntimeError(
                    f"LibreOffice conversion failed (code {proc.returncode}): {err} "
                    f"(attempt {attempt}/{MAX_RETRIES})"
                )
                logger.warning(str(last_error))
                _cleanup_temp(out_dir)
                if attempt < MAX_RETRIES:
                    await asyncio.sleep(RETRY_DELAY)
                continue

            converted_name = doc_path.stem + ".docx"
            converted_path = Path(out_dir) / converted_name

            if not converted_path.exists():
                last_error = RuntimeError(
                    f"LibreOffice did not produce expected output: {converted_path} "
                    f"(attempt {attempt}/{MAX_RETRIES})"
                )
                logger.warning(str(last_error))
                _cleanup_temp(out_dir)
                if attempt < MAX_RETRIES:
                    await asyncio.sleep(RETRY_DELAY)
                continue

            logger.info("Converted .doc -> .docx: %s -> %s (attempt %d)",
                        doc_path.name, converted_path, attempt)
            return converted_path

        except Exception as exc:
            last_error = exc
            logger.warning("LibreOffice conversion error (attempt %d/%d): %s",
                           attempt, MAX_RETRIES, exc)
            _cleanup_temp(out_dir)
            if attempt < MAX_RETRIES:
                await asyncio.sleep(RETRY_DELAY)

    raise last_error or RuntimeError(
        f"LibreOffice conversion failed after {MAX_RETRIES} attempts for {doc_path.name}"
    )


# ── Tier 2: antiword ───────────────────────────────────────────────────────


def parse_with_antiword(path: Path) -> str:
    """Fallback: extract text using antiword with encoding detection."""
    if not is_antiword_available():
        raise RuntimeError("antiword is not installed")

    # Try UTF-8 mapping first (better for Chinese text)
    args = ["antiword", "-w", "0", "-m", "UTF-8.txt", str(path)]
    result = subprocess.run(args, capture_output=True, text=True, timeout=30)

    if result.returncode == 0 and result.stdout.strip():
        text = result.stdout.strip()

        # Detect and fix encoding issues
        try:
            import chardet
            detected = chardet.detect(text.encode("latin-1", errors="replace"))
            if detected and detected.get("encoding") == "UTF-8":
                text = text.encode("latin-1", errors="replace").decode("utf-8", errors="replace")
        except ImportError:
            pass  # chardet not available, use as-is

        return text

    # Fallback: without -m flag
    result = subprocess.run(
        ["antiword", "-w", "0", str(path)],
        capture_output=True, text=True, timeout=30,
    )
    if result.returncode == 0 and result.stdout.strip():
        return result.stdout.strip()

    raise RuntimeError(f"antiword failed: {result.stderr.strip()}")


# ── Tier 3: catdoc ─────────────────────────────────────────────────────────


def parse_with_catdoc(path: Path) -> str:
    """Extract text using catdoc (lightweight, good Chinese support)."""
    if not is_catdoc_available():
        raise RuntimeError("catdoc is not installed")

    # catdoc -w disables line wrapping, -d specifies charset
    result = subprocess.run(
        ["catdoc", "-w", str(path)],
        capture_output=True, text=True, timeout=30,
    )
    if result.returncode == 0 and result.stdout.strip():
        return result.stdout.strip()

    # Try with UTF-8 charset hint
    result = subprocess.run(
        ["catdoc", "-d", "utf-8", str(path)],
        capture_output=True, text=True, timeout=30,
    )
    if result.returncode == 0 and result.stdout.strip():
        return result.stdout.strip()

    raise RuntimeError(f"catdoc failed: {result.stderr.strip()}")


# ── Tier 4: python-docx direct ─────────────────────────────────────────────


def parse_with_docx_direct(path: Path) -> str:
    """Attempt to parse .doc file as if it were actually .docx format.

    Some files have a .doc extension but are actually Office Open XML (.docx).
    """
    try:
        import docx
        doc = docx.Document(str(path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        text = "\n\n".join(parts)
        if len(text) > 50:
            return text
    except Exception:
        pass
    raise RuntimeError("File is not a valid DOCX (tried python-docx direct)")


# ── Tier 5: OLE structured extraction ─────────────────────────────────────


def parse_with_olefile(path: Path) -> str:
    """Last-resort: parse .doc binary via OLE Compound Document structure.

    Extracts the WordDocument stream and 1Table/0Table streams from the OLE
    container, then extracts readable Unicode text using the FIB (File
    Information Block) to locate text positions.
    """
    try:
        import olefile
    except ImportError:
        raise RuntimeError("olefile is not installed. Install with: pip install olefile")

    ole = olefile.OleFileIO(str(path))
    if not ole.exists("WordDocument"):
        ole.close()
        raise RuntimeError("Not a valid OLE Word document (no WordDocument stream)")

    # Read the WordDocument stream
    word_stream = ole.openstream("WordDocument").read()

    # Read FIB (File Information Block) - first 2 bytes of WordDocument
    # give us key structural information
    if len(word_stream) < 2:
        ole.close()
        raise RuntimeError("WordDocument stream too short")

    # Try to extract text using a simpler approach:
    # Read all streams and decode any UTF-16LE text found
    text_parts: list[str] = []

    for stream_name in ole.listdir():
        try:
            stream_data = ole.openstream(stream_name).read()
            # Try UTF-16LE decode (Word's native encoding for text)
            try:
                decoded = stream_data.decode("utf-16-le", errors="ignore")
                # Filter out control characters, keep printable text
                clean = "".join(
                    c for c in decoded
                    if c.isprintable() or c in "\n\r\t" or "一" <= c <= "鿿"
                    or "　" <= c <= "〿"
                )
                if len(clean) > 20:
                    text_parts.append(clean)
            except Exception:
                pass
        except Exception:
            pass

    ole.close()

    if text_parts:
        # Find the longest text part (usually the main document stream)
        text_parts.sort(key=len, reverse=True)
        result = text_parts[0]
        # Remove excessively long runs of nulls/newlines
        import re
        result = re.sub(r"\n{4,}", "\n\n\n", result)
        result = re.sub(r"\0+", "", result)
        return result.strip()

    raise RuntimeError("OLE extraction produced no readable text")


# ── Tier 6: Plain text (binary decode) ────────────────────────────────────


def parse_with_plain_text(path: Path) -> str:
    """Last-resort fallback: extract readable text from .doc binary content.

    Reads the file as bytes, tries multiple encodings, and filters out
    binary noise while keeping readable Chinese and ASCII text.
    """
    raw = path.read_bytes()

    # Try multiple encodings
    text = None
    for encoding in ["utf-8", "utf-16-le", "gbk", "gb2312", "big5"]:
        try:
            decoded = raw.decode(encoding, errors="ignore")
            # Count CJK characters to assess if this encoding is correct
            cjk_count = sum(
                1 for c in decoded
                if "一" <= c <= "鿿" or "　" <= c <= "〿"
            )
            if cjk_count > 10:
                text = decoded
                break
        except Exception:
            continue

    if text is None:
        text = raw.decode("utf-8", errors="replace")

    # Filter out binary garbage
    lines: list[str] = []
    for line in text.splitlines():
        if not line.strip():
            continue

        # Skip lines that are clearly binary garbage
        stripped = line.strip()
        if len(stripped) < 3:
            continue

        # Count printable characters including CJK
        printable = sum(
            1 for c in stripped
            if c.isprintable() or c in "\t"
            or "一" <= c <= "鿿"
            or "　" <= c <= "〿"
            or "＀" <= c <= "￯"
        )
        ratio = printable / max(len(stripped), 1)

        # Keep lines with high printable ratio or significant Chinese text
        cjk_chars = sum(
            1 for c in stripped
            if "一" <= c <= "鿿"
        )
        if ratio > 0.6 or cjk_chars > 5:
            lines.append(stripped)

    if not lines:
        raise RuntimeError("Plain text extraction produced no readable content")

    return "\n".join(lines)


# ── Orchestration ──────────────────────────────────────────────────────────


async def extract_text_from_doc(path: Path) -> str:
    """Extract text from .doc file using the full 6-tier fallback chain.

    Returns the extracted text content.
    """
    file_hash = _compute_file_hash(path)

    # Tier 0: Cache
    cached = _read_cache(file_hash)
    if cached:
        return cached

    errors: list[str] = []

    # Tier 1: LibreOffice → docx → python-docx (best quality)
    if is_libreoffice_available():
        try:
            logger.info("[doc] Tier 1 — LibreOffice conversion: %s", path.name)
            docx_path = await convert_doc_to_docx(path)
            try:
                import docx as _docx
                doc = _docx.Document(str(docx_path))
                parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        parts.append(para.text)
                text = "\n\n".join(parts)
                if text.strip():
                    _write_cache(file_hash, text)
                    return text
            finally:
                try:
                    docx_path.unlink(missing_ok=True)
                    docx_path.parent.rmdir()
                except Exception:
                    pass
        except Exception as e:
            logger.warning("[doc] Tier 1 failed: %s", e)
            errors.append(f"LibreOffice: {e}")

    # Tier 2: antiword + encoding detection
    if is_antiword_available():
        try:
            logger.info("[doc] Tier 2 — antiword extraction: %s", path.name)
            text = parse_with_antiword(path)
            if text.strip():
                _write_cache(file_hash, text)
                return text
        except Exception as e:
            logger.warning("[doc] Tier 2 failed: %s", e)
            errors.append(f"antiword: {e}")

    # Tier 3: catdoc
    if is_catdoc_available():
        try:
            logger.info("[doc] Tier 3 — catdoc extraction: %s", path.name)
            text = parse_with_catdoc(path)
            if text.strip():
                _write_cache(file_hash, text)
                return text
        except Exception as e:
            logger.warning("[doc] Tier 3 failed: %s", e)
            errors.append(f"catdoc: {e}")

    # Tier 4: python-docx direct (file might actually be .docx)
    try:
        logger.info("[doc] Tier 4 — python-docx direct attempt: %s", path.name)
        text = parse_with_docx_direct(path)
        if text.strip():
            _write_cache(file_hash, text)
            return text
    except Exception as e:
        logger.warning("[doc] Tier 4 failed: %s", e)
        errors.append(f"docx-direct: {e}")

    # Tier 5: OLE structured extraction
    try:
        logger.info("[doc] Tier 5 — OLE extraction: %s", path.name)
        text = parse_with_olefile(path)
        if text.strip():
            _write_cache(file_hash, text)
            return text
    except Exception as e:
        logger.warning("[doc] Tier 5 failed: %s", e)
        errors.append(f"olefile: {e}")

    # Tier 6: Plain text (binary decode)
    try:
        logger.info("[doc] Tier 6 — plain text extraction: %s", path.name)
        text = parse_with_plain_text(path)
        if text.strip():
            _write_cache(file_hash, text)
            return text
    except Exception as e:
        logger.warning("[doc] Tier 6 failed: %s", e)
        errors.append(f"plain_text: {e}")

    raise RuntimeError(
        f"All .doc parsing methods failed for {path.name}. "
        f"Errors: {'; '.join(errors)}"
    )


# ── Helpers ────────────────────────────────────────────────────────────────


def _cleanup_temp(dir_path: str) -> None:
    """Best-effort cleanup of a temp directory."""
    try:
        import os as _os
        for f in _os.listdir(dir_path):
            _os.remove(_os.path.join(dir_path, f))
        _os.rmdir(dir_path)
    except Exception:
        pass
