"""DOCX parser using python-docx."""

from pathlib import Path

import docx

from app.services.parsing.types import DocType, ParsedDocument
from app.services.parsing.docling_parser import compute_file_hash


class DocxParser:
    """Parse DOCX files using python-docx."""

    def parse(self, file_path: str | Path, doc_id: str, kb_id: str) -> ParsedDocument:
        path = Path(file_path)
        doc = docx.Document(str(path))

        parts = []
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "")
                prefix = "#" * int(level)
                parts.append(f"{prefix} {para.text}")
            elif para.text.strip():
                parts.append(para.text)

        # Handle tables
        for i, table in enumerate(doc.tables):
            table_md = self._table_to_markdown(table)
            parts.append(f"## Table {i + 1}\n\n{table_md}")

        content = "\n\n".join(parts)

        return ParsedDocument(
            doc_id=doc_id,
            kb_id=kb_id,
            filename=path.name,
            file_type=DocType.DOCX,
            file_hash=compute_file_hash(path),
            content=content,
            metadata={"paragraph_count": len(doc.paragraphs), "source": "python-docx"},
        )

    @staticmethod
    def _table_to_markdown(table) -> str:
        """Convert a python-docx table to Markdown."""
        rows_data = []
        for row in table.rows:
            rows_data.append([cell.text for cell in row.cells])

        if not rows_data:
            return ""

        header = rows_data[0]
        lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(["---"] * len(header)) + " |",
        ]
        for row in rows_data[1:]:
            lines.append("| " + " | ".join(row) + " |")

        return "\n".join(lines)
